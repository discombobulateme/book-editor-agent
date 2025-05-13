import os
import anthropic
from dotenv import load_dotenv
import glob
import time
import re
import argparse
import traceback
import docx
import threading
import signal

# Import our terminal colors utility
from terminal_colors import Colors, print_header, print_subheader, print_stats, success, warning, error, info, debug, Spinner, StatusUpdatingSpinner, ConnectionMonitoringSpinner

# Load environment variables
load_dotenv()

def get_api_key():
    """Get the Anthropic API key from environment variables"""
    api_key = os.getenv("ANTHROPIC_API_KEY")
    if not api_key:
        raise ValueError("No API key found. Please set the ANTHROPIC_API_KEY environment variable.")
    return api_key

def create_anthropic_client(api_key):
    """Create and return an Anthropic client"""
    # Simple initialization without proxy settings for compatibility with version 0.6.0
    return anthropic.Anthropic(api_key=api_key)

def get_text_files(specific_file=None):
    """Get a list of text files (.txt and .docx) from the original-texts directory or a specific file"""
    if specific_file:
        # If a specific file was provided, check if it exists
        if os.path.exists(specific_file):
            return [specific_file]
        # If the specific file doesn't exist, try looking for it in the original-texts directory
        possible_path = os.path.join("original-texts", os.path.basename(specific_file))
        if os.path.exists(possible_path):
            return [possible_path]
        # If still not found, return an empty list
        return []
    
    # Otherwise get all text files in the original-texts directory
    return glob.glob("original-texts/*.txt") + glob.glob("original-texts/*.docx")

def get_review_notes(filename):
    """Get review notes for a text file if they exist"""
    base_name = os.path.basename(filename)
    name, ext = os.path.splitext(base_name)
    
    # Check for a txt review file first
    review_file = os.path.join("review-notes", f"{name}.txt")
    
    if os.path.exists(review_file):
        with open(review_file, "r") as f:
            return f.read()
    
    # If not found, check for a docx review file
    docx_review_file = os.path.join("review-notes", f"{name}.docx")
    if os.path.exists(docx_review_file):
        return read_docx_content(docx_review_file)
    
    return None

def read_file_content(filename):
    """Read the content of a file"""
    if filename.endswith('.docx'):
        return read_docx_content(filename)
    with open(filename, "r") as f:
        return f.read()

def read_docx_content(filename):
    """Read content from a .docx file"""
    doc = docx.Document(filename)
    full_text = []
    for para in doc.paragraphs:
        if para.text:
            full_text.append(para.text)
    return '\n\n'.join(full_text)

def save_edited_text(filename, edited_content, model_name, output_format="same"):
    """Save the edited text to the edited-texts directory with model name in filename"""
    # Extract just the filename from the path
    base_name = os.path.basename(filename)
    
    # Create the directory if it doesn't exist
    os.makedirs("edited-texts", exist_ok=True)
    
    # Split the base name into name and extension
    name, ext = os.path.splitext(base_name)
    
    # Get a shortened model identifier
    model_id = model_name
    
    # Determine output extension based on format preference
    if output_format == "same":
        output_ext = ext
    else:
        output_ext = f".{output_format}"
    
    # Create output path with the determined extension
    output_path = os.path.join("edited-texts", f"{name}-{model_id}{output_ext}")
    
    # Check if the file already exists and add a version number if it does
    version = 1
    while os.path.exists(output_path):
        # Increment version number and create a new filename
        output_path = os.path.join("edited-texts", f"{name}-{model_id}-{version}{output_ext}")
        version += 1
    
    # Write the edited content to the file based on file type
    if output_ext.lower() == '.docx':
        save_docx_content(output_path, edited_content)
    else:
        with open(output_path, "w") as f:
            f.write(edited_content)
    
    info(f"Saved edited text to {output_path}")
    return output_path

def save_docx_content(output_path, content):
    """Save content to a .docx file with formatting"""
    doc = docx.Document()
    
    # Split content into paragraphs and add them to the document
    paragraphs = content.split('\n\n')
    for p in paragraphs:
        if p.strip():
            doc.add_paragraph(p.strip())
    
    doc.save(output_path)

def create_editing_prompt(original_text, review_notes, instructions_content):
    """Create a prompt for the AI to edit the text"""
    prompt = (
        "You are a professional editor skilled in enhancing text without losing content or nuance.\n\n"
        
        "## EDITING INSTRUCTIONS\n\n"
        "1. Edit the following text according to the style guidelines provided below.\n"
        "2. Preserve all important information, facts, and details from the original text.\n"
        "3. Maintain the original text organization and paragraph structure.\n"
        "4. If the review notes request shortening, make the text more concise.\n"
        "5. Otherwise, maintain the full content and approximately the same length.\n"
        "6. Follow the Structured Experiential Theory approach as defined in the style guide.\n"
        "7. Return ONLY the edited text without any explanations or comments.\n\n"
        
        "## STYLE GUIDELINES\n"
        f"{instructions_content}\n\n"
        
        "## ORIGINAL TEXT\n"
        f"{original_text}\n\n"
    )
    
    if review_notes:
        prompt += (
            "## REVIEW NOTES\n"
            f"{review_notes}\n\n"
        )
    
    prompt += "## YOUR EDITED TEXT\n"
    
    return prompt

def is_already_edited(filename, model):
    """Check if a file has already been edited by this model and has no review notes to incorporate"""
    base_name = os.path.basename(filename)
    name, ext = os.path.splitext(base_name)
    
    # Check if an edited version exists (either txt or docx extension)
    edited_glob_txt = os.path.join("edited-texts", f"{name}-{model}*.txt")
    edited_glob_docx = os.path.join("edited-texts", f"{name}-{model}*.docx")
    edited_files = glob.glob(edited_glob_txt) + glob.glob(edited_glob_docx)
    
    # If no edited version exists, the file needs editing
    if not edited_files:
        return False
    
    # If review notes exist, the file should be re-edited
    review_notes = get_review_notes(filename)
    if review_notes:
        return False
    
    # File has been edited and no review notes exist
    return True

def edit_text_with_claude(client, text_file, model, instructions_content, output_format="same", chunk_size=0):
    """Edit a text file using Claude"""
    # Check if the file has already been edited by this model and has no review notes
    if is_already_edited(text_file, model):
        info(f"Skipping {text_file} - already edited by {model} and no review notes found.")
        return None
    
    print_header(f"PROCESSING: {text_file} with model {model}")
    
    # Initialize variables to keep track of resources that need cleanup
    api_spinner = None
    connection_monitor = None
    old_handler = None
    
    try:
        # Read the text file
        spinner = Spinner("Reading input file...").start()
        original_text = read_file_content(text_file)
        spinner.stop(f"Read input file: {text_file}")
        
        # Get review notes if they exist
        spinner = Spinner("Checking for review notes...").start()
        review_notes = get_review_notes(text_file)
        
        if review_notes:
            spinner.stop("Found review notes")
            print_subheader("📝 REVIEW NOTES FOUND")
            info(review_notes)
        else:
            spinner.stop("No review notes found")
            info("📋 No review notes found. Proceeding with standard editing.")
        
        # Get text statistics for original text
        spinner = Spinner("Analyzing text...").start()
        original_word_count = len(original_text.split())
        original_char_count = len(original_text)
        original_paragraphs = len([p for p in original_text.split('\n\n') if p.strip()])
        spinner.stop("Text analysis complete")
        
        print_subheader("📊 ORIGINAL TEXT STATISTICS")
        print_stats("Words", original_word_count)
        print_stats("Characters", original_char_count)
        print_stats("Paragraphs", original_paragraphs)
        
        # Calculate estimated processing time
        estimated_minutes = estimate_processing_time(original_word_count, model)
        
        # Check if this is a large document and provide warning and time estimate
        if original_word_count > 5000:
            if estimated_minutes > 10:
                print_subheader("⚠️ LARGE DOCUMENT DETECTED")
                warning(f"This document has {original_word_count} words.")
                warning(f"Estimated processing time: {estimated_minutes:.1f} minutes.")
                
                if chunk_size == 0:
                    warning(f"Consider using the --chunk-size option for large documents.")
                    warning(f"Example: --chunk-size 5000 will process the document in smaller segments.")
                    
                    # Ask for confirmation before proceeding
                    info("Press Enter to continue, or Ctrl+C to cancel.")
                    try:
                        input()
                    except KeyboardInterrupt:
                        info("Operation cancelled by user.")
                        return None
        
        # If chunking is enabled and document is large enough, process in chunks
        if chunk_size > 0 and original_word_count > chunk_size:
            return process_document_in_chunks(client, text_file, original_text, review_notes, 
                                             instructions_content, model, output_format, chunk_size)
        
        # Create prompt
        spinner = Spinner("Preparing editing prompt...").start()
        prompt = create_editing_prompt(original_text, review_notes, instructions_content)
        
        # Estimate token count and cost
        input_tokens = len(prompt.split()) * 1.3  # Rough estimate: words * 1.3
        spinner.stop(f"Prompt prepared (~{int(input_tokens)} estimated tokens)")
        
        # Set max_tokens based on model
        max_tokens = get_max_tokens_for_model(model)
    
        # Create a connection monitor for real status updates
        connection_monitor = AnthropicConnectionMonitor(client, timeout=300)
        
        # Use a connection monitoring spinner
        api_spinner = ConnectionMonitoringSpinner(
            message=f"Sending request to Claude API ({model})...",
            check_interval=3,
            timeout=360
        ).start_request()
        
        # Start monitoring the connection with estimated tokens
        connection_monitor.start_request(api_spinner, estimated_tokens=input_tokens, model=model)
        
        # Set up a timeout handler
        def timeout_handler(signum, frame):
            if api_spinner:
                api_spinner.stop("Request timed out")
            if connection_monitor:
                connection_monitor.stop_monitoring()
            raise TimeoutError("API request timed out")
        
        # Set a signal alarm for very long operations (use a safe default)
        old_handler = signal.signal(signal.SIGALRM, timeout_handler)
        # Use max(5 minutes, 2x estimated time)
        timeout_seconds = int(max(300, estimated_minutes * 120))
        signal.alarm(timeout_seconds)
        
        try:
            try:
                start_time = time.time()
                message = client.messages.create(
                    model=model,
                    max_tokens=max_tokens,
                    temperature=0.85,
                    messages=[
                        {"role": "user", "content": prompt}
                    ]
                )
                
                # Immediately stop monitoring on successful completion
                end_time = time.time()
                duration = end_time - start_time
                
                # First handle alarms
                signal.alarm(0)  # Cancel any pending alarms immediately
                
                # First stop the connection monitoring - do this BEFORE updating UI
                if connection_monitor:
                    connection_monitor.request_completed = True  # Mark request as completed
                    connection_monitor.stop_monitoring()
                
                # Then stop the spinner with completion message
                if api_spinner:
                    api_spinner.request_in_progress = False  # Mark request as completed
                    api_spinner.stop(f"Request completed in {duration:.1f} seconds")
            except Exception as e:
                # Clean up resources on error
                if connection_monitor:
                    connection_monitor.stop_monitoring()
                if api_spinner:
                    api_spinner.stop(f"Request failed: {str(e)}")
                raise  # Re-raise the exception
            
            # Extract the edited text
            edited_text = message.content[0].text
            
            # Calculate tokens and cost
            input_tokens = message.usage.input_tokens
            output_tokens = message.usage.output_tokens
            total_tokens = input_tokens + output_tokens
            
            # Estimate cost based on model
            cost = estimate_cost(model, input_tokens, output_tokens)
            
            # Display token usage and cost
            print_subheader("💰 API USAGE")
            print_stats("Input tokens", input_tokens)
            print_stats("Output tokens", output_tokens)
            print_stats("Total tokens", total_tokens)
            print_stats("Estimated cost", f"${cost:.4f} USD")
            
            # Clean up the response to remove any metadata
            spinner = Spinner("Cleaning up response...").start()
            edited_text = cleanup_response(edited_text)
            spinner.stop("Response cleaned")
            
            # Get word count stats for edited text
            spinner = Spinner("Analyzing edited text...").start()
            edited_word_count = len(edited_text.split())
            edited_char_count = len(edited_text)
            edited_paragraphs = len([p for p in edited_text.split('\n\n') if p.strip()])
            spinner.stop("Analysis complete")
            
            print_subheader("📊 EDITED TEXT STATISTICS")
            print_stats("Words", edited_word_count, original_word_count)
            print_stats("Characters", edited_char_count, original_char_count)
            print_stats("Paragraphs", edited_paragraphs, original_paragraphs)
            
            # Validate the edited text
            spinner = Spinner("Validating edited text...").start()
            validation_result = validate_edited_text(original_text, edited_text, review_notes)
            
            if not validation_result:
                spinner.stop("Validation failed - text appears to be shortened")
                warning("Attempting to regenerate edited text...")
                
                # Add stronger instructions to prevent shortening
                spinner = Spinner("Preparing retry prompt...").start()
                retrying_prompt = (
                    "You are a professional editor skilled in enhancing text without losing content or nuance.\n\n"
                    
                    "## IMPORTANT CORRECTION NEEDED\n\n"
                    "Your previous edit was too short or appeared to be a summary. Please try again with these requirements:\n"
                    "- Do NOT summarize or condense the text unless specifically asked to in the review notes\n"
                    "- Maintain the FULL length and content of the original text\n"
                    "- Preserve the same number of paragraphs as the original\n"
                    "- Apply the style guidelines while keeping all original details\n\n"
                    
                    "## STYLE GUIDELINES\n"
                    f"{instructions_content}\n\n"
                    
                    "## ORIGINAL TEXT\n"
                    f"{original_text}\n\n"
                )
                
                if review_notes:
                    retrying_prompt += (
                        "## REVIEW NOTES\n"
                        f"{review_notes}\n\n"
                    )
                
                retrying_prompt += "## YOUR CORRECTED EDIT (FULL LENGTH)\n"
                spinner.stop("Retry prompt ready")
                
                # Create a connection monitor for the retry
                retry_connection_monitor = AnthropicConnectionMonitor(client, timeout=180)
                
                # Try again with real status monitoring
                retry_api_spinner = ConnectionMonitoringSpinner(
                    message="Sending retry request to Claude API...",
                    check_interval=3,
                    timeout=240
                ).start_request()
                
                # Start monitoring the retry connection
                retry_estimated_tokens = len(retrying_prompt.split()) * 1.3  # Rough estimate
                retry_connection_monitor.start_request(retry_api_spinner, estimated_tokens=retry_estimated_tokens, model=model)
                
                try:
                    start_time = time.time()
                    message = client.messages.create(
                        model=model,
                        max_tokens=max_tokens,
                        temperature=0.85,
                        messages=[
                            {"role": "user", "content": retrying_prompt}
                        ]
                    )
                    
                    # Calculate time taken
                    duration = time.time() - start_time
                    
                    # Stop the connection monitor
                    if retry_connection_monitor:
                        retry_connection_monitor.stop_monitoring()
                    
                    # Stop the spinner
                    if retry_api_spinner:
                        retry_api_spinner.stop(f"Retry completed in {duration:.1f} seconds")
                except Exception as e:
                    # Ensure monitoring is stopped on error
                    if retry_connection_monitor:
                        retry_connection_monitor.stop_monitoring()
                    if retry_api_spinner:
                        retry_api_spinner.stop(f"Retry failed: {str(e)}")
                    raise  # Re-raise the exception
                
                # Track retry token usage and cost
                retry_input_tokens = message.usage.input_tokens
                retry_output_tokens = message.usage.output_tokens
                total_tokens += (retry_input_tokens + retry_output_tokens)
                
                # Calculate retry cost
                retry_cost = estimate_cost(model, retry_input_tokens, retry_output_tokens)
                cost += retry_cost
                
                # Display updated token usage and cost
                print_subheader("💰 TOTAL API USAGE (INITIAL + RETRY)")
                print_stats("Total tokens", total_tokens)
                print_stats("Estimated cost", f"${cost:.4f} USD")
                
                # Extract the edited text from retry
                edited_text = message.content[0].text
                
                # Clean up the response again
                spinner = Spinner("Cleaning up retry response...").start() 
                edited_text = cleanup_response(edited_text)
                spinner.stop("Retry response cleaned")
                
                # Final validation
                spinner = Spinner("Validating retry result...").start()
                if not validate_edited_text(original_text, edited_text, review_notes):
                    spinner.stop("Validation failed again")
                    warning("AI still produced shortened text. Saving anyway, but please review.")
                else:
                    spinner.stop("Validation successful")
            else:
                spinner.stop("Validation successful")
            
            # Save the edited text
            spinner = Spinner("Saving edited text...").start()
            output_path = save_edited_text(text_file, edited_text, model, output_format)
            spinner.stop(f"Saved to {output_path}")
            
            # Print final statistics
            final_word_count = len(edited_text.split())
            final_char_count = len(edited_text)
            final_paragraphs = len([p for p in edited_text.split('\n\n') if p.strip()])
            
            final_word_ratio = final_word_count / original_word_count * 100
            
            print_subheader("📊 FINAL STATISTICS")
            info(f"Original: {original_word_count} words, {original_paragraphs} paragraphs")
            info(f"Edited:   {final_word_count} words, {final_paragraphs} paragraphs")
            info(f"Ratio:    {final_word_ratio:.1f}% of original length")
            
            success(f"Saved edited text to: {output_path}")
            info(f"{'='*80}\n")
            
            return edited_text
            
        finally:
            # Always cancel the alarm when done with the API call
            signal.alarm(0)
            
            # Restore original signal handler if we set one
            if old_handler:
                try:
                    signal.signal(signal.SIGALRM, old_handler)
                except Exception:
                    pass  # Ignore errors in signal handling restoration
            
            # Stop the connection monitor if it exists
            if connection_monitor:
                try:
                    connection_monitor.stop_monitoring()
                except Exception:
                    pass
    
    except TimeoutError as e:
        error(f"API request timed out: {str(e)}")
        warning("The request to Claude API took too long and was cancelled.")
        warning("Try using the --chunk-size option for large documents.")
        return None
    except Exception as e:
        error(f"Error processing {text_file}: {e}")
        traceback.print_exc()
        return None
    finally:
        # Final resource cleanup
        try:
            if api_spinner:
                api_spinner.stop()
        except Exception:
            pass
        
        try:
            if connection_monitor:
                connection_monitor.stop_monitoring()
        except Exception:
            pass
        
        # Reset signal alarm just to be safe
        signal.alarm(0)

def process_document_in_chunks(client, text_file, original_text, review_notes, instructions_content, model, output_format, chunk_size):
    """Process a large document by breaking it into chunks, editing each, then recombining"""
    print_subheader("🧩 CHUNKING LARGE DOCUMENT")
    
    # Split the document into chunks
    spinner = Spinner("Splitting document into chunks...").start()
    chunks = chunk_document(original_text, max_words=chunk_size)
    spinner.stop(f"Document split into {len(chunks)} chunks")
    
    # Process each chunk
    edited_chunks = []
    total_cost = 0
    total_tokens = 0
    
    for i, chunk in enumerate(chunks):
        chunk_num = i + 1
        print_subheader(f"PROCESSING CHUNK {chunk_num}/{len(chunks)}")
        
        # Create a specific prompt for this chunk
        spinner = Spinner(f"Preparing prompt for chunk {chunk_num}...").start()
        
        # Modify review notes for chunks if needed
        chunk_reviews = None
        if review_notes:
            chunk_reviews = f"CHUNK {chunk_num}/{len(chunks)}: {review_notes}\n\nIMPORTANT: This is chunk {chunk_num} of {len(chunks)}. Focus on editing THIS CHUNK ONLY."
        
        chunk_prompt = create_editing_prompt(chunk, chunk_reviews, instructions_content)
        
        # Estimate token count
        estimated_tokens = len(chunk_prompt.split()) * 1.3
        spinner.stop(f"Chunk prompt prepared (~{int(estimated_tokens)} tokens)")
        
        # Set max_tokens based on model
        max_tokens = get_max_tokens_for_model(model)
        
        # Create a connection monitor
        connection_monitor = AnthropicConnectionMonitor(client, timeout=180)
        
        # Process this chunk with real connection monitoring
        api_spinner = ConnectionMonitoringSpinner(
            message=f"Processing chunk {chunk_num}/{len(chunks)}...",
            check_interval=3,
            timeout=240
        ).start_request()
        
        # Start monitoring the connection
        connection_monitor.start_request(api_spinner, estimated_tokens=estimated_tokens, model=model)
        
        try:
            # Call Claude API
            start_time = time.time()
            message = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                temperature=0.85,
                messages=[
                    {"role": "user", "content": chunk_prompt}
                ]
            )
            
            # Immediately stop monitoring
            processing_time = time.time() - start_time
            
            # Extract the edited chunk
            edited_chunk = message.content[0].text
            
            # Track token usage and cost
            input_tokens = message.usage.input_tokens
            output_tokens = message.usage.output_tokens
            chunk_tokens = input_tokens + output_tokens
            total_tokens += chunk_tokens
            
            # Calculate cost
            chunk_cost = estimate_cost(model, input_tokens, output_tokens)
            total_cost += chunk_cost
            
            # Stop the connection monitor
            if connection_monitor:
                connection_monitor.stop_monitoring()
            
            # Stop the spinner with completion message
            if api_spinner:
                api_spinner.stop(f"Chunk {chunk_num}/{len(chunks)} complete in {processing_time:.1f}s (${chunk_cost:.4f})")
            
            # Clean up the response
            edited_chunk = cleanup_response(edited_chunk)
            
            # Add to our list of edited chunks
            edited_chunks.append(edited_chunk)
            
            # Display progress
            progress_pct = (chunk_num / len(chunks)) * 100
            info(f"Progress: {progress_pct:.1f}% complete")
            info(f"Running cost: ${total_cost:.4f} ({total_tokens} tokens)")
            
        except Exception as e:
            connection_monitor.stop_monitoring()
            api_spinner.stop()
            error(f"Error processing chunk {chunk_num}: {e}")
            # Try to salvage what we have so far
            if edited_chunk := locals().get('edited_chunk'):
                edited_chunks.append(edited_chunk)
                warning(f"Saving partial result for chunk {chunk_num}")
            # Continue with other chunks
    
    # Check if we have any processed chunks
    if not edited_chunks:
        error("Failed to process any chunks successfully")
        return None
    
    # Combine the edited chunks
    spinner = Spinner("Combining edited chunks...").start()
    combined_text = "\n\n".join(edited_chunks)
    spinner.stop("All chunks combined")
    
    # Display combined statistics
    word_count = len(combined_text.split())
    original_word_count = len(original_text.split())
    
    print_subheader("📊 CHUNKED PROCESSING RESULTS")
    print_stats("Original words", original_word_count)
    print_stats("Edited words", word_count)
    print_stats("Word ratio", f"{(word_count/original_word_count*100):.1f}%")
    print_stats("Total chunks", len(chunks))
    print_stats("Processed chunks", len(edited_chunks))
    print_stats("Total tokens", total_tokens)
    print_stats("Total cost", f"${total_cost:.4f}")
    
    # Save the result
    spinner = Spinner("Saving combined result...").start()
    output_path = save_edited_text(text_file, combined_text, f"{model}-chunked", output_format)
    spinner.stop(f"Saved to {output_path}")
    
    success(f"Chunked processing complete! Saved to: {output_path}")
    
    return combined_text

def get_max_tokens_for_model(model):
    """Get appropriate max_tokens value based on model"""
    # Define max tokens for each model version
    max_tokens_map = {
        "claude-3-opus-20240229": 12000,
        "claude-3-sonnet-20240229": 8000,
        "claude-3-haiku-20240307": 4000,
        "claude-3-5-sonnet-20240620": 8000,
        "claude-3-5-haiku-20240620": 4000,
        "claude-3-7-sonnet-20250219": 12000,
        "claude-2.1": 4000,
    }
    
    # Try exact match first
    if model in max_tokens_map:
        return max_tokens_map[model]
    
    # Try to match by model family
    if "3-7" in model:
        return 12000  # Claude 3.7 has large context capacity
    elif "3-5" in model and "haiku" in model:
        return 4000
    elif "3-5" in model and "sonnet" in model:
        return 8000
    elif "haiku" in model:
        return 4000
    elif "sonnet" in model:
        return 8000
    elif "opus" in model:
        return 12000
    elif "2.1" in model:
        return 4000
    
    # Default to a conservative value
    return 4000

def get_available_models():
    """Return a dictionary of available Claude models with their descriptions"""
    return {
        "claude-3-7-sonnet-20250219": "Claude 3.7 Sonnet - Latest high performance model",
        "claude-3-5-sonnet-20240620": "Claude 3.5 Sonnet - Latest balanced model",
        "claude-3-5-haiku-20240620": "Claude 3.5 Haiku - Latest affordable model",
        "claude-3-opus-20240229": "Claude 3 Opus - Highest quality, most expensive",
        "claude-3-sonnet-20240229": "Claude 3 Sonnet - Good balance of quality and cost",
        "claude-3-haiku-20240307": "Claude 3 Haiku - Fastest and most affordable",
        "claude-2.1": "Claude 2.1 - Older model, still reliable"
    }

def sanitize_custom_id(filename):
    """Create a valid custom_id from a filename (alphanumeric, underscore, hyphen, max 64 chars)"""
    # Extract the base name without extension
    base_name = os.path.splitext(os.path.basename(filename))[0]
    # Replace invalid characters with underscores
    sanitized = re.sub(r'[^a-zA-Z0-9_-]', '_', base_name)
    # Truncate to 64 characters if longer
    if len(sanitized) > 64:
        sanitized = sanitized[:64]
    return sanitized

def process_batch_item(client, text_file, model, instructions_content, output_format="same"):
    """Process a single item from a batch, with validation and cleanup"""
    # Read the text file
    spinner = Spinner(f"Reading {text_file}...").start()
    original_text = read_file_content(text_file)
    
    # Get review notes if they exist
    review_notes = get_review_notes(text_file)
    
    if review_notes:
        spinner.stop(f"Read file with review notes: {text_file}")
        info(f"📝 Review notes found")
    else:
        spinner.stop(f"Read file: {text_file}")
    
    # Create prompt
    spinner = Spinner("Preparing prompt...").start()
    prompt = create_editing_prompt(original_text, review_notes, instructions_content)
    spinner.stop("Prompt ready")
    
    # Set max_tokens based on model
    max_tokens = get_max_tokens_for_model(model)
    
    # Track total tokens and cost
    total_tokens = 0
    total_cost = 0

    # Send to Claude
    try:
        # Create a connection monitor
        connection_monitor = AnthropicConnectionMonitor(client, timeout=180)
        
        # Use a connection monitoring spinner
        api_spinner = ConnectionMonitoringSpinner(
            message=f"Sending request to Claude API ({model})...",
            check_interval=3,
            timeout=240
        ).start_request()
        
        # Start monitoring the connection with estimated tokens
        estimated_tokens = len(prompt.split()) * 1.3  # Rough estimate
        connection_monitor.start_request(api_spinner, estimated_tokens=estimated_tokens, model=model)
        
        try:
            start_time = time.time()
            message = client.messages.create(
                model=model,
                max_tokens=max_tokens,
                temperature=0.85,
                messages=[
                    {"role": "user", "content": prompt}
                ]
            )
            
            # Calculate time taken
            duration = time.time() - start_time
            
            # Stop the connection monitor 
            if connection_monitor:
                connection_monitor.stop_monitoring()
            
            # Stop the spinner
            if api_spinner:
                api_spinner.stop(f"Request completed in {duration:.1f} seconds")
                
            # Extract the edited text
            edited_text = message.content[0].text
        except Exception as e:
            # Ensure monitoring is stopped on error
            if connection_monitor:
                connection_monitor.stop_monitoring()
            if api_spinner:
                api_spinner.stop(f"Request failed: {str(e)}")
            raise  # Re-raise the exception
        
        # Track token usage and cost
        input_tokens = message.usage.input_tokens
        output_tokens = message.usage.output_tokens
        total_tokens += (input_tokens + output_tokens)
        
        # Calculate cost
        request_cost = estimate_cost(model, input_tokens, output_tokens)
        total_cost += request_cost
        
        # Display token usage
        info(f"Tokens: {input_tokens} in, {output_tokens} out (${request_cost:.4f})")
        
        # Clean up the response to remove any metadata
        spinner = Spinner("Processing response...").start()
        edited_text = cleanup_response(edited_text)
        spinner.stop("Response processed")
        
        # Validate the edited text
        spinner = Spinner("Validating edit...").start()
        if not validate_edited_text(original_text, edited_text, review_notes):
            spinner.stop("Validation failed - text appears shortened")
            warning(f"Batch item {text_file} produced shortened text - regenerating...")
            
            # Add stronger instructions to prevent shortening
            spinner = Spinner("Preparing retry...").start()
            retrying_prompt = (
                "You are a professional editor skilled in enhancing text without losing content or nuance.\n\n"
                
                "## IMPORTANT CORRECTION NEEDED\n\n"
                "Your previous edit was too short or appeared to be a summary. Please try again with these requirements:\n"
                "- Do NOT summarize or condense the text unless specifically asked to in the review notes\n"
                "- Maintain the FULL length and content of the original text\n"
                "- Preserve the same number of paragraphs as the original\n"
                "- Apply the style guidelines while keeping all original details\n\n"
                
                "## STYLE GUIDELINES\n"
                f"{instructions_content}\n\n"
                
                "## ORIGINAL TEXT\n"
                f"{original_text}\n\n"
            )
            
            if review_notes:
                retrying_prompt += (
                    "## REVIEW NOTES\n"
                    f"{review_notes}\n\n"
                )
            
            retrying_prompt += "## YOUR CORRECTED EDIT (FULL LENGTH)\n"
            spinner.stop("Retry prompt ready")
            
            # Create a connection monitor for the retry
            retry_connection_monitor = AnthropicConnectionMonitor(client, timeout=180)
            
            # Try again with real status monitoring
            retry_api_spinner = ConnectionMonitoringSpinner(
                message="Sending retry request to Claude API...",
                check_interval=3,
                timeout=240
            ).start_request()
            
            # Start monitoring the retry connection
            retry_estimated_tokens = len(retrying_prompt.split()) * 1.3  # Rough estimate
            retry_connection_monitor.start_request(retry_api_spinner, estimated_tokens=retry_estimated_tokens, model=model)
            
            try:
                start_time = time.time()
                message = client.messages.create(
                    model=model,
                    max_tokens=max_tokens,
                    temperature=0.85,
                    messages=[
                        {"role": "user", "content": retrying_prompt}
                    ]
                )
                
                # Calculate time taken
                duration = time.time() - start_time
                
                # Stop the connection monitor
                if retry_connection_monitor:
                    retry_connection_monitor.stop_monitoring()
                
                # Stop the spinner
                if retry_api_spinner:
                    retry_api_spinner.stop(f"Retry completed in {duration:.1f} seconds")
            except Exception as e:
                # Ensure monitoring is stopped on error
                if retry_connection_monitor:
                    retry_connection_monitor.stop_monitoring()
                if retry_api_spinner:
                    retry_api_spinner.stop(f"Retry failed: {str(e)}")
                raise  # Re-raise the exception
            
            # Track retry token usage and cost
            retry_input_tokens = message.usage.input_tokens
            retry_output_tokens = message.usage.output_tokens
            total_tokens += (retry_input_tokens + retry_output_tokens)
            
            # Calculate retry cost
            retry_cost = estimate_cost(model, retry_input_tokens, retry_output_tokens)
            total_cost += retry_cost
            
            # Display token usage
            info(f"Retry tokens: {retry_input_tokens} in, {retry_output_tokens} out (${retry_cost:.4f})")
            
            # Extract the edited text from retry
            edited_text = message.content[0].text
            
            # Clean up the response again
            spinner = Spinner("Processing retry response...").start()
            edited_text = cleanup_response(edited_text)
            spinner.stop("Retry processed")
            
            # Final validation
            spinner = Spinner("Validating retry...").start()
            if not validate_edited_text(original_text, edited_text, review_notes):
                spinner.stop("Final validation failed")
                warning(f"Batch item {text_file} still produced shortened text. Saving anyway.")
            else:
                spinner.stop("Final validation successful")
        else:
            spinner.stop("Validation successful")
        
        # Save the edited text
        spinner = Spinner("Saving edited text...").start()
        output_path = save_edited_text(text_file, edited_text, model, output_format)
        spinner.stop(f"Saved to {output_path}")
        
        success(f"Processed {text_file} - Total cost: ${total_cost:.4f} ({total_tokens} tokens)")
        
        return edited_text
    except Exception as e:
        if 'spinner' in locals():
            spinner.stop()
        error(f"Error processing batch item {text_file}: {e}")
        return None

def batch_edit_texts(client, text_files, model, instructions_content, output_format="same", chunk_size=0):
    """Process multiple text files in a batch request"""
    if not text_files:
        warning("No text files found in original-texts directory.")
        return
    
    # Track batch totals
    batch_total_tokens = 0
    batch_total_cost = 0
    processed_count = 0
    skipped_count = 0
    
    # Filter files to only those with review notes
    spinner = Spinner("Checking for files with review notes...").start()
    files_with_notes = []
    for text_file in text_files:
        if get_review_notes(text_file):
            files_with_notes.append(text_file)
        else:
            skipped_count += 1
    
    if not files_with_notes:
        spinner.stop("No files with review notes found")
        warning("No files with review notes found. Add review notes to process files.")
        return
    
    spinner.stop(f"Found {len(files_with_notes)} files with review notes (skipped {skipped_count})")
    
    # Display large document warnings if applicable
    large_docs = []
    for text_file in files_with_notes:
        # Check file size
        spinner = Spinner(f"Analyzing {os.path.basename(text_file)}...").start()
        try:
            content = read_file_content(text_file)
            word_count = len(content.split())
            
            if word_count > 5000:
                large_docs.append((text_file, word_count))
            
            spinner.stop()
        except Exception as e:
            spinner.stop(f"Error reading {text_file}")
            error(f"Could not analyze {text_file}: {str(e)}")
    
    # Warn about large documents
    if large_docs:
        print_subheader("⚠️ LARGE DOCUMENTS DETECTED")
        for doc, words in large_docs:
            est_time = estimate_processing_time(words, model)
            warning(f"{os.path.basename(doc)}: {words} words (~{est_time:.1f} min)")
        
        if chunk_size == 0:
            warning("Large documents detected. Consider using --chunk-size for better processing.")
            warning("Example: --chunk-size 5000 will process large documents in smaller segments.")
            
            # Ask for confirmation
            info("Press Enter to continue with batch processing, or Ctrl+C to cancel.")
            try:
                input()
            except KeyboardInterrupt:
                info("Batch operation cancelled by user.")
                return
    
    # Process files individually
    success(f"Processing {len(files_with_notes)} files...")
    
    for i, text_file in enumerate(files_with_notes):
        print_subheader(f"BATCH ITEM {i+1}/{len(files_with_notes)}: {text_file}")
        
        # Record token count before processing
        pre_tokens = get_total_tokens_used(client, model)
        
        # Process the item, using chunking if enabled
        spinner = Spinner(f"Processing {os.path.basename(text_file)}...").start()
        start_time = time.time()
        
        result = edit_text_with_claude(client, text_file, model, instructions_content, output_format, chunk_size)
        
        processing_time = time.time() - start_time
        spinner.stop(f"Processed in {processing_time:.1f} seconds")
        
        # Get tokens used in this operation
        post_tokens = get_total_tokens_used(client, model)
        tokens_this_file = post_tokens - pre_tokens
        
        # Estimate cost for this file
        # Assuming average token cost for simplicity
        if "opus" in model:
            avg_token_cost = 0.00004
        elif "sonnet" in model:
            avg_token_cost = 0.00001
        else:  # haiku or others
            avg_token_cost = 0.0000007
        
        cost_this_file = tokens_this_file * avg_token_cost
        
        # Update batch totals
        if result:
            batch_total_tokens += tokens_this_file
            batch_total_cost += cost_this_file
            processed_count += 1
        
        # Show progress
        progress_pct = ((i + 1) / len(files_with_notes)) * 100
        info(f"Batch progress: {progress_pct:.1f}% ({i+1}/{len(files_with_notes)} files)")
        info(f"Running cost: ${batch_total_cost:.4f} ({batch_total_tokens} tokens)")
        
    # Final summary
    print_subheader("🔶 BATCH PROCESSING SUMMARY")
    print_stats("Files processed", processed_count)
    print_stats("Files skipped", skipped_count)
    print_stats("Total tokens used", batch_total_tokens)
    print_stats("Estimated total cost", f"${batch_total_cost:.4f} USD")
    
    success("Batch processing complete!")

def get_total_tokens_used(client, model):
    """Placeholder function to track token usage
    In a real implementation, this could use Anthropic's reporting API"""
    # This is just a placeholder - Claude API doesn't currently 
    # have a built-in way to track cumulative token usage
    # Future implementation could use client.get_usage() if available
    return 0  # For now we rely on per-request tracking

def validate_edited_text(original_text, edited_text, review_notes=None):
    """Validate that the edited text is not significantly shorter than the original
    unless specifically requested in review notes.
    Returns True if the text is valid, False if it needs to be redone"""
    
    # Calculate basic statistics
    original_words = len(original_text.split())
    edited_words = len(edited_text.split())
    word_ratio = edited_words / original_words
    
    # Count paragraphs
    original_paragraphs = len([p for p in original_text.split('\n\n') if p.strip()])
    edited_paragraphs = len([p for p in edited_text.split('\n\n') if p.strip()])
    
    # Print statistics
    info(f"Original text: {original_words} words, {original_paragraphs} paragraphs")
    info(f"Edited text: {edited_words} words, {edited_paragraphs} paragraphs")
    info(f"Word ratio: {word_ratio*100:.1f}%")
    
    # Check for summary indicators
    summary_phrases = [
        "the text below", "this text", "this is a", "below is a", 
        "condensed version", "shorter version", "summary of"
    ]
    
    # Check if the edited text starts with a summary indicator
    first_100_words = " ".join(edited_text.split()[:100]).lower()
    has_summary_indicator = any(phrase in first_100_words for phrase in summary_phrases)
    
    # If review notes exist, we allow some shortening
    if review_notes:
        # Only flag as invalid if very heavily shortened (less than 50%)
        if word_ratio < 0.5:
            warning(f"Edited text is excessively shortened: {word_ratio*100:.1f}% of original length")
            return False
        
        # Accept the edit if it has a reasonable length
        success(f"Edited text length acceptable: {word_ratio*100:.1f}% of original")
        return True
    
    # If no review notes exist (no shortening should happen):
    
    # Check if too short (less than 90% of original)
    if word_ratio < 0.9:
        warning(f"Edited text is too short ({word_ratio*100:.1f}% of original)")
        return False
    
    # Check for significant paragraph structure changes
    if edited_paragraphs < original_paragraphs * 0.9 and original_paragraphs > 3:
        warning(f"Edited text has fewer paragraphs ({edited_paragraphs} vs {original_paragraphs})")
        return False
    
    # Check for summary indicators when no shortening was requested
    if has_summary_indicator:
        warning(f"Edited text appears to be a summary rather than an edit")
        return False
        
    success(f"Validation successful: {edited_words} words ({word_ratio*100:.1f}% of original)")
    return True

def cleanup_response(text):
    """Clean up the response from the model by removing any metadata or notes at the end"""
    
    # Common endings that models might add
    ending_phrases = [
        "This text has been edited",
        "I have edited the text",
        "The edited text follows",
        "Here is the edited text",
        "I've maintained the full length",
        "I've preserved all content",
        "This edit maintains",
        "Edited according to",
        "Following the style guidelines",
        "As per the instructions",
    ]
    
    # Check for any of the ending phrases and remove them and anything that follows
    cleaned_text = text
    for phrase in ending_phrases:
        if phrase in cleaned_text:
            # Split on the phrase and take only what comes before it
            cleaned_text = cleaned_text.split(phrase)[0].strip()
    
    # Also remove anything after common markdown or comment delimiters if they appear near the end
    ending_delimiters = ["---", "***", "###", "```", "//"]
    for delimiter in ending_delimiters:
        # Only consider delimiters in the last 15% of the text to avoid removing content
        search_start = int(len(cleaned_text) * 0.85)
        last_part = cleaned_text[search_start:]
        if delimiter in last_part:
            # Find the position in the full text
            delimiter_pos = cleaned_text.rfind(delimiter, search_start)
            # Check if there's text after this delimiter that looks like metadata
            text_after = cleaned_text[delimiter_pos:].lower()
            if any(phrase.lower() in text_after for phrase in ["edit", "note", "comment", "text", "follow"]):
                cleaned_text = cleaned_text[:delimiter_pos].strip()
    
    return cleaned_text

def estimate_cost(model, input_tokens, output_tokens):
    """Estimate the cost of API usage based on model and tokens"""
    # Claude pricing as of July 2024 (subject to change)
    pricing = {
        "claude-3-opus-20240229": {"input": 15.0, "output": 75.0},
        "claude-3-sonnet-20240229": {"input": 3.0, "output": 15.0},
        "claude-3-haiku-20240307": {"input": 0.25, "output": 1.25},
        "claude-3-5-sonnet-20240620": {"input": 3.0, "output": 15.0},
        "claude-3-5-haiku-20240620": {"input": 0.25, "output": 1.25},
        "claude-3-7-sonnet-20250219": {"input": 5.0, "output": 15.0},
        "claude-2.1": {"input": 0.8, "output": 2.4},
    }
    
    # Find the right pricing tier
    model_pricing = None
    
    # Try exact match first
    if model in pricing:
        model_pricing = pricing[model]
    else:
        # Try partial match based on model family
        if "3-7" in model and "sonnet" in model:
            model_pricing = pricing["claude-3-7-sonnet-20250219"]
        elif "3-5" in model and "sonnet" in model:
            model_pricing = pricing["claude-3-5-sonnet-20240620"]
        elif "3-5" in model and "haiku" in model:
            model_pricing = pricing["claude-3-5-haiku-20240620"]
        elif "opus" in model:
            model_pricing = pricing["claude-3-opus-20240229"]
        elif "sonnet" in model:
            model_pricing = pricing["claude-3-sonnet-20240229"]
        elif "haiku" in model:
            model_pricing = pricing["claude-3-haiku-20240307"]
        elif "2.1" in model:
            model_pricing = pricing["claude-2.1"]
    
    # If model not found, use haiku pricing (lowest) as fallback
    if not model_pricing:
        model_pricing = pricing["claude-3-haiku-20240307"]
    
    # Calculate cost (price is per million tokens, so divide by 1,000,000)
    input_cost = (input_tokens * model_pricing["input"]) / 1000000
    output_cost = (output_tokens * model_pricing["output"]) / 1000000
    
    return input_cost + output_cost

def chunk_document(text, max_words=5000, preserve_paragraphs=True):
    """
    Split a large document into manageable chunks.
    
    Args:
        text (str): The text to split into chunks
        max_words (int): Maximum number of words per chunk
        preserve_paragraphs (bool): Whether to preserve paragraph boundaries
        
    Returns:
        list: List of text chunks
    """
    # If the text is small enough, return it as a single chunk
    words = text.split()
    if len(words) <= max_words:
        return [text]
    
    chunks = []
    
    if preserve_paragraphs:
        # Split by paragraphs (double newlines)
        paragraphs = [p for p in text.split("\n\n") if p.strip()]
        
        current_chunk = []
        current_word_count = 0
        
        for paragraph in paragraphs:
            paragraph_words = len(paragraph.split())
            
            # If adding this paragraph would exceed the limit, start a new chunk
            if current_word_count + paragraph_words > max_words and current_chunk:
                # Join the current chunk and add it to the list
                chunks.append("\n\n".join(current_chunk))
                current_chunk = [paragraph]
                current_word_count = paragraph_words
            else:
                # Add this paragraph to the current chunk
                current_chunk.append(paragraph)
                current_word_count += paragraph_words
        
        # Add the last chunk if it exists
        if current_chunk:
            chunks.append("\n\n".join(current_chunk))
    else:
        # Simple chunking by word count
        for i in range(0, len(words), max_words):
            chunk = " ".join(words[i:i+max_words])
            chunks.append(chunk)
    
    return chunks

def estimate_processing_time(word_count, model):
    """
    Estimate how long processing will take based on word count and model.
    
    Args:
        word_count (int): Number of words in the document
        model (str): Name of the model
    
    Returns:
        float: Estimated processing time in minutes
    """
    # Base processing rates (words per second) for different models
    # These are rough estimates and may need adjustment
    processing_speeds = {
        "claude-3-opus-20240229": 500,     # Claude 3 Opus - most thorough, slowest
        "claude-3-sonnet-20240229": 800,   # Claude 3 Sonnet - balanced
        "claude-3-haiku-20240307": 1200,   # Claude 3 Haiku - fastest
        "claude-3-5-sonnet-20240620": 850, # Claude 3.5 Sonnet - slightly faster than 3
        "claude-3-5-haiku-20240620": 1300, # Claude 3.5 Haiku - fastest
        "claude-3-7-sonnet-20250219": 900, # Claude 3.7 Sonnet - high performance
        "claude-2.1": 700,                 # Claude 2.1 - older model
    }
    
    # Default to Claude 3 Sonnet rate if model not recognized
    rate = 800  # Default processing rate
    
    # Try exact match first
    if model in processing_speeds:
        rate = processing_speeds[model]
    else:
        # Try partial match based on model family
        if "3-7" in model and "sonnet" in model:
            rate = processing_speeds["claude-3-7-sonnet-20250219"]
        elif "3-5" in model and "sonnet" in model:
            rate = processing_speeds["claude-3-5-sonnet-20240620"]
        elif "3-5" in model and "haiku" in model:
            rate = processing_speeds["claude-3-5-haiku-20240620"]
        elif "opus" in model:
            rate = processing_speeds["claude-3-opus-20240229"]
        elif "sonnet" in model:
            rate = processing_speeds["claude-3-sonnet-20240229"]
        elif "haiku" in model:
            rate = processing_speeds["claude-3-haiku-20240307"]
        elif "2.1" in model:
            rate = processing_speeds["claude-2.1"]
    
    # Calculate estimated time in minutes
    # Add a base overhead time for API setup, etc.
    estimated_seconds = (word_count / rate) + 5
    return estimated_seconds / 60

class AnthropicConnectionMonitor:
    """Monitors the connection status to Anthropic API and provides heartbeat checks"""

    def __init__(self, client, timeout=120):
        """
        Initialize the connection monitor
        
        Args:
            client: Anthropic client instance
            timeout: Maximum seconds to wait before considering a connection timed out
        """
        self.client = client
        self.timeout = timeout
        self.last_heartbeat = None
        self.heartbeat_thread = None
        self.stop_heartbeat = threading.Event()
        self.connection_status = "Initializing"
        self.lock = threading.Lock()
        self.request_in_progress = False
        self.request_completed = False  # Flag to indicate request has completed
        self.request_start_time = None
        self.expected_duration = None
        self.estimated_tokens = 0

    def start_request(self, spinner, estimated_tokens=0, model=""):
        """Start monitoring a request"""
        self.request_in_progress = True
        self.request_completed = False  # Reset completion flag
        self.request_start_time = time.time()
        self.last_heartbeat = time.time()
        
        # Estimate expected duration based on token count and model
        self.estimated_tokens = estimated_tokens
        words_per_second = 800  # Default to Sonnet processing rate
        
        # Adjust based on model
        if "opus" in model.lower():
            words_per_second = 500
        elif "3-7" in model.lower() and "sonnet" in model.lower():
            words_per_second = 900
        elif "3-5" in model.lower() and "sonnet" in model.lower():
            words_per_second = 850
        elif "3-5" in model.lower() and "haiku" in model.lower():
            words_per_second = 1300
        elif "sonnet" in model.lower():
            words_per_second = 800
        elif "haiku" in model.lower():
            words_per_second = 1200
        elif "2.1" in model.lower():
            words_per_second = 700
            
        # Add base overhead (API setup, etc.)
        self.expected_duration = (estimated_tokens / 1.3) / words_per_second + 5
        
        # Set connection status
        self.connection_status = "Request started"
        
        # Update the spinner with initial estimates
        if spinner:
            spinner.connection_object = self
            spinner.update_activity()
        
        # Start the heartbeat thread to periodically update status
        self.stop_heartbeat.clear()
        self.heartbeat_thread = threading.Thread(target=self._heartbeat_loop, args=(spinner,))
        self.heartbeat_thread.daemon = True
        self.heartbeat_thread.start()
        
        return self

    def _heartbeat_loop(self, spinner):
        """Background thread that updates connection status"""
        while not self.stop_heartbeat.is_set():
            # Check if the request is explicitly marked as completed
            if self.request_completed:
                # Stop monitoring automatically if request is completed
                self.stop_heartbeat.set()
                self.request_in_progress = False
                break
                
            if self.request_in_progress:
                # Calculate elapsed time
                elapsed_time = time.time() - self.request_start_time
                
                # Update status based on elapsed time vs expected duration
                with self.lock:
                    if elapsed_time > self.timeout:
                        self.connection_status = f"Connection may be stalled (timeout exceeded)"
                    elif elapsed_time > self.expected_duration * 1.5:
                        self.connection_status = f"Taking longer than expected ({elapsed_time:.0f}s / ~{self.expected_duration:.0f}s)"
                    elif elapsed_time > self.expected_duration:
                        self.connection_status = f"Almost complete ({int(elapsed_time/self.expected_duration*90)}%)"
                    else:
                        progress = min(95, int(elapsed_time/self.expected_duration*100))
                        self.connection_status = f"Processing (~{progress}% complete)"
                
                # Update spinner if provided
                if spinner:
                    spinner.update_activity()
                    
                # Set last heartbeat time
                self.last_heartbeat = time.time()
                
            # Sleep for a while before updating again
            time.sleep(2)

    def update_activity(self):
        """Update the last activity timestamp"""
        self.last_heartbeat = time.time()
        
    def check_status(self):
        """Get the current connection status"""
        with self.lock:
            if not self.request_in_progress:
                return "Idle"
            
            elapsed_time = time.time() - self.request_start_time
            time_since_heartbeat = time.time() - self.last_heartbeat
            
            # Check for long time since last heartbeat
            if time_since_heartbeat > 30:
                return f"Connection may be lost (no heartbeat for {int(time_since_heartbeat)}s)"
                
            # Return the current connection status
            return self.connection_status
    
    def stop_monitoring(self):
        """Stop the heartbeat thread and cleanup"""
        self.request_in_progress = False
        self.request_completed = True  # Mark the request as explicitly completed
        self.stop_heartbeat.set()
        
        if self.heartbeat_thread and self.heartbeat_thread.is_alive():
            try:
                self.heartbeat_thread.join(timeout=1.0)
            except Exception:
                pass  # Ignore any issues with thread joining
            
        return self

def main():
    # Set up argument parser
    parser = argparse.ArgumentParser(description="Book Editor Agent using Claude AI")
    parser.add_argument("--model", "-m", 
                        default="claude-3-haiku-20240307",
                        help="Claude model to use for editing")
    parser.add_argument("--list-models", "-l", action="store_true",
                        help="List available Claude models with descriptions")
    parser.add_argument("--batch", "-b", action="store_true",
                        help="Process files in batch mode")
    parser.add_argument("--output-format", "-o", 
                        choices=["txt", "docx", "same"],
                        default="same",
                        help="Output format for edited files (txt, docx, or same as input)")
    parser.add_argument("--chunk-size", type=int, default=0,
                        help="Split large documents into chunks of this many words (0 disables chunking)")
    parser.add_argument("--only-with-notes", action="store_true",
                        help="Only process files that have review notes")
    parser.add_argument("file", nargs="?", help="Specific file to process (optional)")
    
    args = parser.parse_args()
    
    # List models if requested
    if args.list_models:
        print_header("AVAILABLE CLAUDE MODELS")
        models = get_available_models()
        # Display models with their pricing
        print_subheader("Model | Input Price | Output Price | Description")
        print_subheader("------|-------------|-------------|------------")
        for model, description in models.items():
            # Get pricing information
            if "opus" in model:
                input_price, output_price = "$15.00", "$75.00"
            elif "3-7" in model and "sonnet" in model:
                input_price, output_price = "$5.00", "$15.00"
            elif "3-5" in model and "sonnet" in model:
                input_price, output_price = "$3.00", "$15.00"
            elif "sonnet" in model:
                input_price, output_price = "$3.00", "$15.00"
            elif "haiku" in model:
                input_price, output_price = "$0.25", "$1.25"
            elif "2.1" in model:
                input_price, output_price = "$0.80", "$2.40"
            else:
                input_price, output_price = "Unknown", "Unknown"
                
            info(f"{model} | {input_price}/M | {output_price}/M | {description}")
        
        info("\nPrices are per million tokens (M). Subject to change - see Anthropic pricing page for latest.")
        return
    
    # Get API key and create client
    spinner = Spinner("Connecting to Anthropic API...").start()
    try:
        api_key = get_api_key()
        client = create_anthropic_client(api_key)
        spinner.stop("Connected to Anthropic API")
    except ValueError as e:
        spinner.stop("API key error")
        error(str(e))
        info("Make sure you have set the ANTHROPIC_API_KEY environment variable.")
        return
    except Exception as e:
        spinner.stop("Connection error")
        error(f"Failed to initialize Anthropic client: {str(e)}")
        return
    
    # Get text files
    spinner = Spinner("Finding text files...").start()
    text_files = get_text_files(args.file)
    
    if not text_files:
        spinner.stop("No files found")
        warning("No text files (.txt or .docx) found in original-texts directory.")
        return
    
    spinner.stop(f"Found {len(text_files)} text files")
    info(f"Output format: {args.output_format}")
    
    # Read instructions for style guidelines
    spinner = Spinner("Loading style guidelines...").start()
    try:
        instructions_content = read_file_content("INSTRUCTIONS.md")
        spinner.stop("Successfully loaded style guidelines")
    except FileNotFoundError:
        spinner.stop("Instructions file not found")
        warning("INSTRUCTIONS.md not found. Using default style guidelines.")
        instructions_content = "Default style guidelines: Academic yet accessible writing with clear explanations."
    
    # Show cost estimate for selected model
    print_subheader("💰 COST INFORMATION")
    if "opus" in args.model:
        info(f"Model: {args.model} (Premium tier - most expensive)")
        info(f"Input pricing: $15.00 per million tokens")
        info(f"Output pricing: $75.00 per million tokens")
    elif "3-7" in args.model and "sonnet" in args.model:
        info(f"Model: {args.model} (High-performance tier)")
        info(f"Input pricing: $5.00 per million tokens")
        info(f"Output pricing: $15.00 per million tokens")
    elif "3-5" in args.model and "sonnet" in args.model:
        info(f"Model: {args.model} (Standard tier)")
        info(f"Input pricing: $3.00 per million tokens")
        info(f"Output pricing: $15.00 per million tokens")
    elif "sonnet" in args.model:
        info(f"Model: {args.model} (Standard tier)")
        info(f"Input pricing: $3.00 per million tokens")
        info(f"Output pricing: $15.00 per million tokens")
    elif "haiku" in args.model:
        info(f"Model: {args.model} (Economy tier - most affordable)")
        info(f"Input pricing: $0.25 per million tokens")
        info(f"Output pricing: $1.25 per million tokens")
    
    # Process files
    if args.batch:
        # Process files in batch
        print_header(f"BATCH PROCESSING {len(text_files)} FILES WITH {args.model}")
        batch_edit_texts(client, text_files, args.model, instructions_content, args.output_format, args.chunk_size)
    else:
        # Process files individually
        files_processed = 0
        for text_file in text_files:
            # Only process files that have review notes
            review_notes = get_review_notes(text_file)
            if review_notes:
                edit_text_with_claude(client, text_file, args.model, instructions_content, args.output_format, args.chunk_size)
                files_processed += 1
            else:
                info(f"Skipping {text_file} - no review notes found.")
        
        if files_processed == 0:
            warning("No files with review notes were found. Add review notes to process files.")
        else:
            success(f"Processed {files_processed} files with review notes.")
            
    success("Book editing process complete!")

if __name__ == "__main__":
    main() 